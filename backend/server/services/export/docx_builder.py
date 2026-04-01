import io
import re
from typing import Optional, List
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from server.config import settings


TITLE_CANDIDATE_RE = re.compile(r"^[\u4e00-\u9fffA-Za-z0-9（）()《》【】、·\s]{2,20}$")
SECTION_RE = re.compile(r"^(一|二|三|四|五|六|七|八|九|十|[0-9]+)[、.]")
ARTICLE_RE = re.compile(r"^第[一二三四五六七八九十0-9]+条")


def _mark_uncertain(line: str, conf: Optional[float]) -> str:
    if conf is None or conf >= settings.low_confidence_threshold:
        return line
    if "：" in line:
        label = line.split("：", 1)[0]
        return f"{label}：【待确认】"
    if ":" in line:
        label = line.split(":", 1)[0]
        return f"{label}:【待确认】"
    return "【待确认】"


def _extract_lines(results: List[dict]) -> List[str]:
    sorted_results = sorted(results, key=lambda r: r.get("order", 0))
    all_lines: list[str] = []
    for item in sorted_results:
        lines = item.get("lines") or []
        if lines:
            for line in lines:
                text = str(line.get("text", "")).strip()
                if not text:
                    continue
                conf = line.get("conf")
                all_lines.append(_mark_uncertain(text, conf))
        else:
            text = str(item.get("text", "")).strip()
            if text:
                for raw in text.splitlines():
                    if raw.strip():
                        all_lines.append(raw.strip())
        all_lines.append("")
    return all_lines


def _guess_title(lines: List[str]) -> List[str]:
    candidates: List[str] = []
    for line in lines:
        clean = line.strip()
        if not clean:
            continue
        if "：" in clean or ":" in clean:
            break
        if SECTION_RE.match(clean) or ARTICLE_RE.match(clean):
            break
        if TITLE_CANDIDATE_RE.match(clean):
            candidates.append(clean)
        if len(candidates) >= 2:
            break
    return candidates


def build_docx(results: list[dict]) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    lines = _extract_lines(results)
    title_lines = _guess_title(lines)
    if title_lines:
        for line in title_lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        doc.add_paragraph("")
        lines = lines[len(title_lines):]
    else:
        title = datetime.now().strftime("合同识别结果_%Y%m%d_%H%M")
        p = doc.add_paragraph(title)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        doc.add_paragraph("")

    for line in lines:
        if not line.strip():
            doc.add_paragraph("")
            continue
        para = doc.add_paragraph(line)
        para.paragraph_format.first_line_indent = Pt(24)
        text = line.strip()
        if SECTION_RE.match(text) or ARTICLE_RE.match(text):
            para.paragraph_format.first_line_indent = Pt(0)
            if para.runs:
                para.runs[0].bold = True

    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
